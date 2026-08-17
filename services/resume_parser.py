import re

from pathlib import Path

from pypdf import PdfReader
from docx import Document


class ResumeParserError(Exception):
    """Raised when a resume cannot be parsed."""


class ResumeParser:

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".txt",
    }

    def extract_text(self, file_path):
        """
        Extract raw text from PDF, DOCX or TXT.
        """

        path = Path(file_path)

        if not path.exists():
            raise ResumeParserError(
                "Resume file was not found."
            )

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ResumeParserError(
                "Unsupported resume format."
            )

        try:

            if extension == ".pdf":
                text = self._extract_pdf(path)

            elif extension == ".docx":
                text = self._extract_docx(path)

            elif extension == ".txt":
                text = self._extract_txt(path)

            else:
                text = ""

        except ResumeParserError:
            raise

        except Exception as exc:

            raise ResumeParserError(
                f"Unable to extract resume text: {exc}"
            ) from exc

        text = self.clean_text(text)

        if not text:
            raise ResumeParserError(
                "No readable text was found in the resume."
            )

        return text

    def _extract_pdf(self, path):

        reader = PdfReader(
            str(path)
        )

        pages = []

        for page in reader.pages:

            page_text = page.extract_text()

            if page_text:
                pages.append(
                    page_text
                )

        return "\n".join(
            pages
        )

    def _extract_docx(self, path):

        document = Document(
            str(path)
        )

        paragraphs = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(
                    text
                )

        # Also extract text from tables.
        for table in document.tables:

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    cell_text = cell.text.strip()

                    if cell_text:
                        cells.append(
                            cell_text
                        )

                if cells:
                    paragraphs.append(
                        " | ".join(cells)
                    )

        return "\n".join(
            paragraphs
        )

    def _extract_txt(self, path):

        encodings = [
            "utf-8",
            "utf-8-sig",
            "cp1252",
            "latin-1",
        ]

        for encoding in encodings:

            try:

                return path.read_text(
                    encoding=encoding
                )

            except UnicodeDecodeError:
                continue

        raise ResumeParserError(
            "Unable to decode the text file."
        )

    def clean_text(self, text):

        if not text:
            return ""

        # Normalize line endings.
        text = text.replace(
            "\r\n",
            "\n"
        )

        text = text.replace(
            "\r",
            "\n"
        )

        # Remove null characters.
        text = text.replace(
            "\x00",
            ""
        )

        # Normalize spaces.
        text = re.sub(
            r"[ \t]+",
            " ",
            text
        )

        # Remove excessive blank lines.
        text = re.sub(
            r"\n{3,}",
            "\n\n",
            text
        )

        # Remove spaces at beginning/end of lines.
        lines = []

        for line in text.splitlines():

            line = line.strip()

            if line:
                lines.append(
                    line
                )

        return "\n".join(
            lines
        ).strip()


resume_parser = ResumeParser()